import configparser
import logging
import os
from logging import exception

from docx import Document
from docx.shared import Inches


from configurazioni import configurazione
from estrai_info import estraiInfo

MyConf = configurazione()
dataConfLogging=MyConf.recupera_loggin_config()
logging.getLogger(__name__)
logging.basicConfig(filename=str(dataConfLogging["logginfile"]),filemode=dataConfLogging['filemode'],level=logging.DEBUG,
                        format='%(asctime)s - %(levelname)s - %(message)s')
def main():
    try:
        logging.info("chiamata WriteReport")
        writeReport()
        logging.info("fine chiamata WriteReport")
    except Exception as e:
        logging.error(e)

def writeReport():
    logging.info("WriteReport")
    config = configparser.ConfigParser()
    config.read("config.ini")
    Threshold = MyConf.recupera_Threshold()
    estrazioneinfo = estraiInfo(logging)
    nomeTemplateFile,Nome_report_da_salvare,nomeRepo,resultNmeaFile=MyConf.ritorno_Configurazioni()
    listNmeaDirecotry = os.listdir(nomeRepo)
    TableStyle = MyConf.ritorno_elenco_config("TableStyle")

   #region list of images from config file
    lista_immagini_da_ricercare_SatAvOS = MyConf.recupera_immagini_daConfig("ImmaginiSatAvOS")
    lista_immagini_da_ricercare_SatAvPRS = MyConf.recupera_immagini_daConfig("ImmaginiSatAvPRS")
    lista_immagini_da_ricercare_SatGeoOS = MyConf.recupera_immagini_daConfig("ImmaginiGeometryOS")
    lista_immagini_da_ricercare_SatGeoPRS = MyConf.recupera_immagini_daConfig("ImmaginiGeometryPRS")
    lista_immagini_da_ricercare_Horiz_PosOS=MyConf.recupera_immagini_daConfig("HorizImgResPosOS")
    lista_immagini_da_ricercare_Horiz_PosPRS = MyConf.recupera_immagini_daConfig("HorizImgResPosPRS")
    lista_immagini_da_ricercare_Vert_PosOS=MyConf.recupera_immagini_daConfig("VertImgResPosOS")
    lista_immagini_da_ricercare_Vert_PosPRS=MyConf.recupera_immagini_daConfig("VertImgResPosPRS")
    lista_immagini_da_ricercare_Plot=MyConf.recupera_immagini_daConfig("PositionImageENU")
   #endregion

    document = Document(nomeTemplateFile)
    subfolders = [f.path for f in os.scandir(nomeRepo) if f.is_dir()]
    listapulitasubfolder = estraiInfo.pulisci_listaNome(subfolders)
    #region immagini estratte dalla folder e scremate.
    result_imgae_finded_satellite_AvailableOS = estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_SatAvOS)
    result_imgae_finded_satellite_AvailablePRS = estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,
                                                                                          lista_immagini_da_ricercare_SatAvPRS)
    result_image_finded_sat_geoOS = estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,
                                                                               lista_immagini_da_ricercare_SatGeoOS)
    result_image_finded_sat_geoPRS = estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_SatGeoPRS)

    result_image_finded_ReultPos_HorOS= estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_Horiz_PosOS)
    result_image_finded_ReultPos_HorPRS= estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_Horiz_PosPRS)
    result_image_finded_PlotEnu=estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_Plot)
    result_image_finded_ReultPos_VertOS= estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_Vert_PosOS)
    result_image_finded_ReultPos_VertPRS= estrazioneinfo.recupera_tutte_le_immagini(nomeRepo,lista_immagini_da_ricercare_Vert_PosPRS)
    #endregion immagini estratte dalla folder e scremate.
    result = estraiInfo.recupera_Tutti_file_risultati_Test_Nmea(resultNmeaFile, nomeRepo)

    for folder in listapulitasubfolder:
        dati_da_mergiare = []
        for filedaleggere in result:
            print(folder)
            if folder in filedaleggere:
                dateTimeTest = folder.rsplit("_", 1)[1][:-6]
                print(dateTimeTest)
                dati_da_mergiare.append(filedaleggere)
                print(filedaleggere)
        dati_estratti = estrazioneinfo.estrai_dati(dati_da_mergiare)

    #region  DATA OBSERVATION
        try:
            logging.info("Data Observation")
            document.add_heading(config.get("data_observation", "Chapt_name") + " " + dateTimeTest, int(config.get("data_observation","level")))
            recupero_header_TabellaDataObs = MyConf.ritorno_elenco_config("TableDataObs")
            tableDataObs = document.add_table(rows=1, cols=len(recupero_header_TabellaDataObs))
            tableDataObs.style = TableStyle[0]
            tableDataObs.autofit = True
            tableDataObs.border = True
            hdr_cells = tableDataObs.rows[0].cells
            for j, value in enumerate(recupero_header_TabellaDataObs):
                hdr_cells[j].text = value
            typeSignal = "OS"
            row_cells = tableDataObs.add_row().cells
            print("Tabella dei tempi ")
            row_cells[0].text = "Intervallo1" #+ str(j+1)
            row_cells[1].text= "FROM"+str(estrazioneinfo.TOWsToDateTime(dati_estratti['TOW iniziale'+typeSignal]))+"\nTO:"+str(estrazioneinfo.TOWsToDateTime(dati_estratti['TOW finale'+typeSignal]))
            row_cells[2].text = "FROM:" + str(dati_estratti['TOW iniziale'+typeSignal]) + "\n TO:" + str(dati_estratti['TOW finale'+typeSignal])
        except Exception as ex:
            logging.error("Exception:"+ex)
    #endregion  DATA OBSERVATION

    #region  TEST SATELLITE AVAILABILITY
        try:
            logging.info("satellite availability")
            document.add_heading(config.get("TEST_satellite_availability", "Chapt_name"), int(config.get("TEST_satellite_availability","level")))
            print(folder)
            #region Image TEST SATELLITE AVAILABILITY
            logging.info("Image TEST SATELLITE AVAILABILITY")
            for img_satellite_AvailablePRS in result_imgae_finded_satellite_AvailablePRS:
                for immagine_da_caricare in lista_immagini_da_ricercare_SatAvPRS:
                     if folder+'_PRS' in img_satellite_AvailablePRS:
                         if immagine_da_caricare in img_satellite_AvailablePRS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailablePRS, width=Inches(5.0))
                            run.add_text("")
                            if "E1A" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_satellite_availability","E1A"))
                            if "E6A" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_satellite_availability", "E6A"))

            for img_satellite_AvailableOS in result_imgae_finded_satellite_AvailableOS:
                for immagine_da_caricare in lista_immagini_da_ricercare_SatAvOS:
                    if folder + '_OS' in img_satellite_AvailableOS:
                        if immagine_da_caricare in img_satellite_AvailableOS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailableOS, width=Inches(5.0))
                            document.add_paragraph('')
                            run.add_text( config.get("Didascalie_satellite_availability", "E1BC"))
        except Exception as ex:
            logging.error("Exception image :"+ex)
        #endregion Image TEST SATELLITE AVAILABILITY

        #region Tabella TEST SATELLITE AVAILABILITY
        try:
            logging.info("Tabella TEST SATELLITE AVAILABILITY")
            paraSat = document.add_paragraph().add_run("Table6[" + dateTimeTest + "]:" + config.get("TextTableAv", "Table6"))

            recupero_header_TabellaSatAv = MyConf.ritorno_elenco_config("TableSatAv")
            tableSatAv = document.add_table(rows=1, cols=len(recupero_header_TabellaSatAv))
            tableSatAv.style = TableStyle[0]  #'Table Grid'
            tableSatAv.autofit = True
            tableSatAv.border = True
            hdr_cellsSatAv = tableSatAv.rows[0].cells
            for i, value in enumerate(recupero_header_TabellaSatAv):
                hdr_cellsSatAv[i].text = value

            list_subfield_sat = MyConf.ritorno_elenco_config("SatelliteSubfileds")
            count_Intervall_SatAv=0
            for value in list_subfield_sat:
                row_cellsSatAv = tableSatAv.add_row().cells
                count=0
                if count_Intervall_SatAv==0:
                    row_cellsSatAv[0].text = ("intervallo1")
                    count_Intervall_SatAv+=1
                row_cellsSatAv[1].text = value
                if "Avarage number" in value:
                    row_cellsSatAv[2].text = str(dati_estratti['MediaSatVistaE1APRS'])  # + typeSignal])
                    row_cellsSatAv[3].text = str(dati_estratti['MediaSatVistaE6APRS'])  # + typeSignal])
                    row_cellsSatAv[4].text = str(dati_estratti['MediaSatVistaE1BCOS'])  # + typeSignal])
                elif "Minimum number" in value:
                    row_cellsSatAv[2].text = str(dati_estratti['MinSatInVistaE1APRS'])
                    row_cellsSatAv[3].text = str(dati_estratti['MinSatInVistaE6APRS'])
                    row_cellsSatAv[4].text = str(dati_estratti['MinSatInVistaE1BCOS'])
                elif "Maximum avarage" in value:
                    row_cellsSatAv[2].text = str(dati_estratti['MaxSatInVistaE1APRS'])
                    row_cellsSatAv[3].text = str(dati_estratti['MaxSatInVistaE6APRS'])
                    row_cellsSatAv[4].text = str(dati_estratti['MaxSatInVistaE1BCOS'])
        except Exception as ex:
            logging.error("Exception:"+ex)
        # endregion Tabella TEST SATELLITE AVAILABILITY
    # endregion  TEST SATELLITE AVAILABILITY

    # region Satellite_geometry
        #region image Satellite Geometry
        try:
            logging.info("image Satellite Geometry")
            document.add_heading("5.5.2" + config.get("Satellite_geometry", "Chapt_name"), 1)
            for img_satellite_AvailablePRS in result_image_finded_sat_geoPRS:
                for immagine_da_caricare in lista_immagini_da_ricercare_SatGeoPRS:
                    if folder+'_PRS' in img_satellite_AvailablePRS:
                        if immagine_da_caricare in img_satellite_AvailablePRS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailablePRS, width=Inches(5.0))
                            if "SUT Distribution" in immagine_da_caricare:
                                run.add_text('\n' + config.get("Didascalie_satellite_geometry", "PdopDist"))
                            else:
                                run.add_text('\n' + config.get("Didascalie_satellite_geometry", "Pdop"))

            for img_satellite_AvailableOS in result_image_finded_sat_geoOS:
                 for immagine_da_caricare in lista_immagini_da_ricercare_SatGeoOS:
                     if folder + '_OS' in img_satellite_AvailableOS:
                         if immagine_da_caricare in img_satellite_AvailableOS:
                              paragraph = document.add_paragraph()
                              run = paragraph.add_run()
                              run.add_picture(img_satellite_AvailableOS, width=Inches(5.0))
                              if "SUT Distribution" in immagine_da_caricare:
                                  run.add_text( config.get("Didascalie_satellite_geometryOS", "PdopDist"))
                              else:
                                  run.add_text( config.get("Didascalie_satellite_geometryOS", "Pdop"))
        except Exception as ex:
            logging.error("Exception image :" + ex)
        # endregion image Satellite Geometry

        #region tabella Satellite Geometry
        try:
            logging.info("tabella Satellite Geometry")
            paraSG = document.add_paragraph().add_run(
            "Table7[" + dateTimeTest + "]:" + config.get("TextTableGeometry", "Table7"))
            recupero_header_TabellaSG = MyConf.ritorno_elenco_config("TableSG")
            tableSatGeo = document.add_table(rows=1, cols=len(recupero_header_TabellaSG))
            tableSatGeo.style = TableStyle[0]
            tableSatGeo.autofit = True
            tableSatGeo.border = True
            hdr_cellsSatGeo = tableSatGeo.rows[0].cells
            count_hdrSatGEo = 0
            for i, value in enumerate(recupero_header_TabellaSG):
                hdr_cellsSatGeo[i].text = value
            list_subfield_geometry = MyConf.ritorno_elenco_config("MetricSubfileds")

            for value in list_subfield_geometry:
                row_cellsSatGeo = tableSatGeo.add_row().cells
                if count_hdrSatGEo == 0:
                    row_cellsSatGeo[0].text = ("intervallo1")
                    count_hdrSatGEo += 1
                row_cellsSatGeo[1].text = value
                if "Avarage PDOP" in value:
                    row_cellsSatGeo[2].text = str(dati_estratti['Media PDOPOS'])
                    row_cellsSatGeo[3].text = str(dati_estratti['Media PDOPPRS'])

                elif "Minimum PDOP" in value:
                    row_cellsSatGeo[2].text = str(dati_estratti['Min PDOPOS'])
                    row_cellsSatGeo[3].text = str(dati_estratti['Min PDOPPRS'])

                elif "Maximum PDOP" in value:
                    row_cellsSatGeo[2].text = str(dati_estratti['Max PDOPOS'])
                    row_cellsSatGeo[3].text =str(dati_estratti['Max PDOPPRS'])

                elif "Threshold PDOP" in value:
                    row_cellsSatGeo[2].text = Threshold
                    row_cellsSatGeo[3].text = Threshold
        except Exception as ex:
            logging.error("Exception image :" + ex)

        #endregion tabella Satellite Geometry
    # endregion Satellite:geometry

    #region Observation
        try:
            logging.info("tabella  Observation")
            document.add_heading(config.get("Observation", "OBChapt_name"), int(config.get("Observation", "level")))
            paraObs = document.add_paragraph().add_run("Table8[" + dateTimeTest + "]:" + config.get("TextTableOBS", "Table8"))
            recupero_header_TabellaObs = MyConf.ritorno_elenco_config("TableObs")
            tableObs = document.add_table(rows=1, cols=len(recupero_header_TabellaObs))
            tableObs.style = TableStyle[0]
            tableObs.autofit = True
            tableObs.border = True
            hdr_cellsObs = tableObs.rows[0].cells
            for i, value in enumerate(recupero_header_TabellaObs):
                hdr_cellsObs[i].text = value
            row_cellsObs = tableObs.add_row().cells
            row_cellsObs[0].text ="intevallo1"
            row_cellsObs[1].text = str(dati_estratti['disponibilityOS'])
            row_cellsObs[2].text = str(dati_estratti['disponibilityPRS'])
        except Exception as ex:
            logging.error("Exception image :" + ex)
    #endregion Observation

    #region Resulting Position Accurancy
        #region immagini Result Position Accurancy
        try:
            logging.info("tabella  immagini Result Position Accurancy")
            document.add_heading( config.get("ResultingPosition", "ResPosChap_name"), 1)
            document.add_heading( config.get("ResultingPosition", "SubHorResPosChap_name"), 2)
        #region horizontal
            for img_satellite_AvailablePRS in result_image_finded_ReultPos_HorPRS:
                for immagine_da_caricare in lista_immagini_da_ricercare_Horiz_PosPRS:
                    if folder+"_PRS" in img_satellite_AvailablePRS:
                        if immagine_da_caricare in img_satellite_AvailablePRS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailablePRS, width=Inches(5.0))
                            if"Error vs PDOP" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Horizzontal","ImgHor1"))
                            elif "Horizontal Error vs Percentile.png" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Horizzontal", "ImgHor2"))
            for img_satellite_AvailableOS in result_image_finded_ReultPos_HorOS:
                for immagine_da_caricare in lista_immagini_da_ricercare_Horiz_PosOS:
                    if folder+'_OS' in img_satellite_AvailableOS:
                        if immagine_da_caricare in img_satellite_AvailableOS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailableOS, width=Inches(5.0))
                            if "Error vs PDOP" in immagine_da_caricare:
                               run.add_text(config.get("Didascalie_Horizzontal", "ImgHor3"))
                            elif "Horizontal Error vs Percentile.png" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Horizzontal", "ImgHor4"))
        except Exception as ex:
            logging.error("Exception image :" + ex)
        #endregion horizontal
        #region Plot ENU
        try:
            logging.info("immagini Plot ENU")
            for img_satellite_Plot in result_image_finded_PlotEnu:
                for immagine_da_caricarePlot in lista_immagini_da_ricercare_Plot:
                    if folder+"_PRS" in img_satellite_Plot:
                        if immagine_da_caricarePlot in img_satellite_Plot:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_Plot, width=Inches(5.0))
                            if "PositionENUvsPercentile.png" in immagine_da_caricarePlot:
                                 run.add_text(config.get("Didascalie_ENU", "EnuPercentilePRS"))
                            elif "Position PLOT ENU.png" in immagine_da_caricarePlot:
                                 run.add_text(config.get("Didascalie_ENU", "PlotENUPRS"))
                    if folder+"_OS" in img_satellite_Plot:
                        if immagine_da_caricarePlot in img_satellite_Plot:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_Plot, width=Inches(5.0))
                            if "PositionENUvsPercentile.png" in immagine_da_caricarePlot:
                                run.add_text(config.get("Didascalie_ENU", "EnuPercentileOS"))
                            elif "Position PLOT ENU.png" in immagine_da_caricarePlot:
                                run.add_text(config.get("Didascalie_ENU", "PlotENUOS"))
        except Exception as ex:
            logging.error("Exception image :" + ex)
        #endregion PlotENU
        #region Vertical
        try:
            logging.info("immagini Vertical Position")
            document.add_heading(config.get("ResultingPosition", "SubVertResPosChap_name"), 2)
            for img_satellite_AvailablePRS in result_image_finded_ReultPos_VertPRS:
                for immagine_da_caricare in lista_immagini_da_ricercare_Vert_PosPRS:
                    if folder+"_PRS" in img_satellite_AvailablePRS:
                        if immagine_da_caricare in img_satellite_AvailablePRS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailablePRS, width=Inches(5.0))
                            if "Error vs PDOP" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Vertical", "ImgVer1"))
                            elif "Vertical Error vs Percentile.png" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Vertical", "ImgVer2"))
            for img_satellite_AvailableOS in result_image_finded_ReultPos_VertOS:
                for immagine_da_caricare in lista_immagini_da_ricercare_Vert_PosOS:
                    if folder+"_OS" in img_satellite_AvailableOS:
                        if immagine_da_caricare in img_satellite_AvailableOS:
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_satellite_AvailableOS, width=Inches(5.0))
                            if "Error vs PDOP" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Vertical", "ImgVer3"))
                            elif "Vertical Error vs Percentile.png" in immagine_da_caricare:
                                run.add_text(config.get("Didascalie_Vertical", "ImgVer4"))
        except Exception as ex:
            logging.error("Exception image :" + ex)
        #endregion Vertical
        #endregion immagini Result Position Accurancy
    #region Table Result Position Accurancy
        try:
            logging.info("Table Result Position Accurancy")
            document.add_heading(config.get("ResultingPosition", "SubSumResPosChap_name"),
                             level=int(config.get("ResultingPosition", "level")))
            para = document.add_paragraph().add_run("Table9[" + dateTimeTest + "]:" + config.get("TextTable", "Table9"))
            recupero_header_tabella_pos_accurancy = MyConf.ritorno_elenco_config("TableSummary")
            table_summary = document.add_table(rows=1, cols=len(recupero_header_tabella_pos_accurancy))
            table_summary.style = TableStyle[0]
            table_summary.autofit = True
            table_summary.border = True
            hdr_cells_summary = table_summary.rows[0].cells
            count_summary_intervall = 0

            for i,value in enumerate(recupero_header_tabella_pos_accurancy):
                hdr_cells_summary[i].text=value
            list_subfield_summary=MyConf.ritorno_elenco_config("SubFieldsPositionError")
            for value in list_subfield_summary:
                row_cells_summary = table_summary.add_row().cells
                row_cells_summary[1].text=value
                if count_summary_intervall==0:
                    row_cells_summary[0].text="Intervallo1"
                    count_summary_intervall+=1
                #region HPE
                if "HPE Avarage" in value:
                    row_cells_summary[2].text=str(dati_estratti['Media errore 2DOS'])
                    row_cells_summary[3].text = str(dati_estratti['Media errore 2DPRS'])
                elif"HPE 95 percentile" in value:
                    row_cells_summary[2].text = str(dati_estratti['Percentile 2DPDOPOS'])
                    row_cells_summary[3].text = str(dati_estratti['Percentile 2DPDOPPRS'])
                elif "HPE Horizontal Bias" in value:
                    row_cells_summary[2].text = str(dati_estratti['bias2DOS'])
                    row_cells_summary[3].text = str(dati_estratti['bias2DPRS'])
                #endregion HPE

            #region VPE
                elif "VPE Avarage" in value:
                    row_cells_summary[2].text = str(dati_estratti['Media errore HOS'])
                    row_cells_summary[3].text = str(dati_estratti['Media errore HPRS'])
                elif "VPE 95th percentile" in value:
                    row_cells_summary[2].text = str(dati_estratti['Percentile HOS'])
                    row_cells_summary[3].text = str(dati_estratti['Percentile HPRS'])
                elif "VPE Vertical Bias" in value:
                    row_cells_summary[2].text = str(dati_estratti['bias HOS'])
                    row_cells_summary[3].text = str(dati_estratti['bias HPRS'])
                #endregion VPE

            #region 3D
                elif "3D err Avarage" in value:
                    row_cells_summary[2].text = str(dati_estratti['Media errore 3DOS'])
                    row_cells_summary[3].text = str(dati_estratti['Media errore 3DPRS'])
                elif "3D err percentile" in value:
                    row_cells_summary[2].text = str(dati_estratti['Percentile 3DOS'])
                    row_cells_summary[3].text = str(dati_estratti['Percentile 3DPRS'])
                elif "3D err Horizontal Bias" in value:
                    row_cells_summary[2].text = str(dati_estratti['bias 3DOS'])
                    row_cells_summary[3].text = str(dati_estratti['bias 3DPRS'])
            #endregion 3D
        except Exception as ex:
            logging.error("Exception image :" + ex)
        #endregion Table Result Position Accurancy
        # endregion Resulting Position Accurancy



        document.add_page_break()

    document.save(Nome_report_da_salvare)



if __name__ == "__main__":
    main()
